import os
import re
from typing import List, Dict

DOCX_AVAILABLE = False
PDF_AVAILABLE = False

try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    print("Warning: python-docx not installed. Text parser will be used as a fallback for .docx files.")

try:
    import pdfplumber
    PDF_AVAILABLE = True
except ImportError:
    print("Warning: pdfplumber not installed. Text parser will be used as a fallback for .pdf files.")

class DocumentParser:
    @staticmethod
    def parse_docx(file_path: str) -> List[Dict]:
        """
        Parses a DOCX file and extracts paragraphs.
        """
        if not DOCX_AVAILABLE:
            return DocumentParser.parse_fallback_txt(file_path)
            
        segments = []
        try:
            doc = docx.Document(file_path)
            position = 0
            for idx, para in enumerate(doc.paragraphs):
                text = para.text.strip()
                if not text:
                    continue
                
                # Determine type based on style or length
                seg_type = "body"
                if para.style.name.startswith("Heading"):
                    seg_type = "heading"
                elif para.style.name.startswith("List"):
                    seg_type = "list_item"
                elif len(text) < 60 and not text.endswith((".", "?", "!")):
                    seg_type = "heading" # Simple heuristic
                
                segments.append({
                    "text": text,
                    "type": seg_type,
                    "position": position
                })
                position += 1
                
            # Parse Tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text and not any(s["text"] == cell_text for s in segments):
                            segments.append({
                                "text": cell_text,
                                "type": "table_cell",
                                "position": position
                            })
                            position += 1
                            
        except Exception as e:
            print(f"python-docx failed: {e}. Switching to fallback parsing.")
            return DocumentParser.parse_fallback_txt(file_path)
            
        return segments

    @staticmethod
    def parse_pdf(file_path: str) -> List[Dict]:
        """
        Parses a PDF file page by page and splits content into sentences.
        """
        if not PDF_AVAILABLE:
            return DocumentParser.parse_fallback_txt(file_path)
            
        segments = []
        try:
            position = 0
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if not text:
                        continue
                    
                    # Split into sentences using a simple regex split
                    sentences = re.split(r'(?<=[.!?])\s+', text)
                    for sent in sentences:
                        sent_clean = sent.strip().replace('\n', ' ')
                        if len(sent_clean) < 3:
                            continue
                        
                        seg_type = "body"
                        if len(sent_clean) < 40 and not sent_clean.endswith((".", "?", "!")):
                            seg_type = "heading"
                            
                        segments.append({
                            "text": sent_clean,
                            "type": seg_type,
                            "position": position
                        })
                        position += 1
        except Exception as e:
            print(f"pdfplumber failed: {e}. Switching to fallback parsing.")
            return DocumentParser.parse_fallback_txt(file_path)
            
        return segments

    @staticmethod
    def parse_fallback_txt(file_path: str) -> List[Dict]:
        """
        Fallback parser when document libraries are not available or files are corrupt.
        Reads raw binary file content and decodes what it can.
        """
        segments = []
        position = 0
        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                lines = f.readlines()
                for line in lines:
                    line_clean = line.strip()
                    if len(line_clean) < 3:
                        continue
                    segments.append({
                        "text": line_clean,
                        "type": "body",
                        "position": position
                    })
                    position += 1
        except Exception as e:
            print(f"Fallback parser failed: {e}")
        return segments

    @classmethod
    def parse_document(cls, file_path: str) -> List[Dict]:
        ext = os.path.splitext(file_path)[1].lower()
        if ext == ".docx":
            return cls.parse_docx(file_path)
        elif ext == ".pdf":
            return cls.parse_pdf(file_path)
        else:
            return cls.parse_fallback_txt(file_path)
